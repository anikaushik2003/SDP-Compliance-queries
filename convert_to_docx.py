"""Convert design-document.md to design-document.docx using python-docx."""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, "design-document.md")
DOCX_PATH = os.path.join(SCRIPT_DIR, "design-document.docx")
IMG_PATH = os.path.join(SCRIPT_DIR, "SDP Cumulative Architecture Plan.png")

def parse_md(text):
    """Parse markdown into a list of (type, content) blocks."""
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            blocks.append(("hr", ""))
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            blocks.append(("heading", (level, m.group(2).strip())))
            i += 1
            continue

        # Image
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", line)
        if m:
            blocks.append(("image", (m.group(1), m.group(2))))
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(("code", "\n".join(code_lines)))
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-|: ]+\|", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(re.sub(r"^[-*]\s+", "", lines[i]))
                i += 1
            blocks.append(("bullet", items))
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append(("numbered", items))
            continue

        # Paragraph
        if line.strip():
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,4}|[-*]\s|\d+\.\s|```|\||---|\!\[)", lines[i]):
                para_lines.append(lines[i])
                i += 1
            blocks.append(("para", " ".join(para_lines)))
            continue

        i += 1

    return blocks


def add_formatted_text(paragraph, text):
    """Add text with basic inline formatting (**bold**, `code`)."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        else:
            paragraph.add_run(part)


def parse_table_row(line):
    """Parse a markdown table row into cells."""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def build_docx(blocks):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for btype, content in blocks:
        if btype == "hr":
            # Add a thin line paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            continue

        if btype == "heading":
            level, text = content
            heading_map = {1: 0, 2: 1, 3: 2, 4: 3}
            h = doc.add_heading(level=heading_map.get(level, 2))
            add_formatted_text(h, text)

        elif btype == "image":
            alt, path = content
            full_path = os.path.join(SCRIPT_DIR, path.replace("%20", " "))
            if os.path.exists(full_path):
                doc.add_picture(full_path, width=Inches(6))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(f"[Image: {alt}]")
                p.italic = True

        elif btype == "code":
            for code_line in content.split("\n"):
                p = doc.add_paragraph(code_line, style="No Spacing")
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)

        elif btype == "table":
            # First row is header, second is separator, rest are data
            header = parse_table_row(content[0])
            data_rows = [parse_table_row(r) for r in content[2:]]  # skip separator
            ncols = len(header)

            table = doc.add_table(rows=1 + len(data_rows), cols=ncols, style="Light Grid Accent 1")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            for j, cell_text in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(10)

            # Data
            for i, row_data in enumerate(data_rows):
                for j, cell_text in enumerate(row_data):
                    if j < ncols:
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_formatted_text(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(10)

        elif btype == "bullet":
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_text(p, item)

        elif btype == "numbered":
            for item in content:
                p = doc.add_paragraph(style="List Number")
                add_formatted_text(p, item)

        elif btype == "para":
            p = doc.add_paragraph()
            add_formatted_text(p, content)

    doc.save(DOCX_PATH)
    print(f"Created: {DOCX_PATH}")


if __name__ == "__main__":
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()
    blocks = parse_md(md_text)
    build_docx(blocks)
